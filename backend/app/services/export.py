from __future__ import annotations
"""
Builds NRB statistical reports in Excel, PDF, Word, and CSV formats.

Each report follows the format used by NRB:
  - Region header row (highlighted)
  - County data rows: NPR (M/F/T) | OTHERS (M/F/T) | GRAND TOTAL
  - Region subtotal row
  - Grand total row at the bottom
"""
import csv
import io
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak,
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
# Shared across every NRB report (general + Usajili Mashinani) for a uniform
# look: plain bold/no-fill headers, RED for the row that introduces a group
# (region or county heading), YELLOW for the row that closes it with a sum
# (TOTAL / GRAND TOTAL) — matching NRB's own stats.xlsx convention rather
# than an invented color scheme.
_NRB_FONT   = "Times New Roman"
_NRB_RED    = "FF0000"
_NRB_YELLOW = "FFFF00"

# ── Module definitions ────────────────────────────────────────────────────────
NRB_CATS = ("npr", "replacements", "changes", "duplicates", "type4", "type5")
OTHERS_CATS = ("replacements", "changes", "duplicates", "type4", "type5")

MODULES = [
    ("app", "APPLICATIONS SENT TO HEADQUARTERS"),
    ("ids", "IDs RECEIVED FROM HEADQUARTERS"),
    ("rej", "REJECTIONS FROM HEADQUARTERS"),
]

# ── Data helpers ──────────────────────────────────────────────────────────────

def _sum_others(row: dict, prefix: str) -> tuple[int, int]:
    """Return (male, female) sum for all non-NPR categories."""
    m = sum(row.get(f"{prefix}_{c}_male", 0) for c in OTHERS_CATS)
    f = sum(row.get(f"{prefix}_{c}_female", 0) for c in OTHERS_CATS)
    return m, f


def build_region_county_data(
    submissions: list,
    station_lookup: dict,  # {station_id: Station}
    year: int,
    month: "int | None" = None,
) -> dict:
    """
    Aggregate approved submissions into a nested dict:
      {module_prefix: {region: {county: {npr_m, npr_f, oth_m, oth_f, ...}}}}
    """
    # Initialise structure
    data: dict[str, dict[str, dict[str, dict[str, int]]]] = {}
    for prefix, _ in MODULES:
        data[prefix] = {}

    for sub in submissions:
        station = station_lookup.get(sub.station_id)
        if not station:
            continue
        region = station.region
        county = station.county

        for prefix, _ in MODULES:
            data[prefix].setdefault(region, {}).setdefault(county, {
                "npr_m": 0, "npr_f": 0,
                "oth_m": 0, "oth_f": 0,
            })
            d = data[prefix][region][county]
            d["npr_m"] += getattr(sub, f"{prefix}_npr_male", 0)
            d["npr_f"] += getattr(sub, f"{prefix}_npr_female", 0)
            om, of_ = _sum_others(sub.__dict__, prefix)
            d["oth_m"] += om
            d["oth_f"] += of_

    return data


def _period_label(year: int, month: "int | None") -> str:
    months = ["JAN","FEB","MAR","APR","MAY","JUN","JUL","AUG","SEP","OCT","NOV","DEC"]
    if month:
        return f"{months[month - 1]} {year}"
    return f"JAN – DEC {year}"


# ── Shared table-data builder ─────────────────────────────────────────────────

def _module_table_rows(
    region_data: dict[str, dict[str, dict[str, int]]],
) -> list[list]:
    """
    Returns rows ready to write into any format:
      Each entry is a dict with _type ('region'|'county'|'subtotal'|'grand') + values.
    """
    rows = []
    g = {"npr_m": 0, "npr_f": 0, "oth_m": 0, "oth_f": 0}

    for region in sorted(region_data):
        rows.append({"_type": "region", "label": f"{region} REGION"})
        rt = {"npr_m": 0, "npr_f": 0, "oth_m": 0, "oth_f": 0}

        for county in sorted(region_data[region]):
            d = region_data[region][county]
            rows.append({
                "_type": "county",
                "label": f"{county} COUNTY",
                **d,
            })
            for k in rt:
                rt[k] += d[k]
                g[k]  += d[k]

        rt["_type"]  = "subtotal"
        rt["label"]  = f"{region} REGION TOTALS"
        rows.append(rt)

    g["_type"] = "grand"
    g["label"] = "GRAND TOTAL"
    rows.append(g)
    return rows


def _general_narrative(data: dict, period: str) -> list[str]:
    """Plain-English summary, one paragraph per module — the same figures
    as the tables, in words, for when a written summary is needed."""
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731
    lines = []
    for prefix, module_title in MODULES:
        rows = _module_table_rows(data.get(prefix, {}))
        grand = next((r for r in rows if r["_type"] == "grand"), None)
        if not grand:
            continue
        npr_total = grand["npr_m"] + grand["npr_f"]
        oth_total = grand["oth_m"] + grand["oth_f"]
        total = npr_total + oth_total
        if total == 0:
            lines.append(f"No {module_title.lower()} were recorded for {period}.")
            continue
        male_total = grand["npr_m"] + grand["oth_m"]
        female_total = grand["npr_f"] + grand["oth_f"]
        sentence = (
            f"For {module_title.lower()} during {period}, a total of {total:,} were recorded, "
            f"comprising {npr_total:,} ({pct(npr_total, total):.1f}%) NPR and {oth_total:,} "
            f"({pct(oth_total, total):.1f}%) other categories (replacements, changes, duplicates, "
            f"type 4 and type 5). Of these, {male_total:,} ({pct(male_total, total):.1f}%) were male "
            f"and {female_total:,} ({pct(female_total, total):.1f}%) were female."
        )
        regions = [r for r in rows if r["_type"] == "subtotal"]
        if regions:
            ranked = sorted(regions, key=lambda r: r["npr_m"] + r["npr_f"] + r["oth_m"] + r["oth_f"], reverse=True)
            top = ranked[0]
            top_total = top["npr_m"] + top["npr_f"] + top["oth_m"] + top["oth_f"]
            top_name = top["label"].replace(" REGION TOTALS", "")
            sentence += f" {top_name} Region recorded the highest volume at {top_total:,}."
        lines.append(sentence)
    return lines


# ── Excel ──────────────────────────────────────────────────────────────────────

def _xl_cell_style(ws, row, col, value, fill=None, font=None, align="center", border=None):
    c = ws.cell(row=row, column=col, value=value)
    if fill:
        c.fill = fill
    if font:
        c.font = font
    if border:
        c.border = border
    c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
    return c


def build_excel_report(
    title: str,
    year: int,
    month: "int | None",
    data: dict,
) -> bytes:
    wb = Workbook()
    wb.remove(wb.active)  # remove default sheet

    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    red_fill    = PatternFill(start_color=_NRB_RED,    end_color=_NRB_RED,    fill_type="solid")
    yellow_fill = PatternFill(start_color=_NRB_YELLOW, end_color=_NRB_YELLOW, fill_type="solid")

    title_font     = Font(name=_NRB_FONT, bold=True, size=14)
    subtitle_font  = Font(name=_NRB_FONT, bold=True, size=12)
    narrative_hdr_font = Font(name=_NRB_FONT, bold=True, size=13)
    narrative_font = Font(name=_NRB_FONT, size=12)
    hdr_font     = Font(name=_NRB_FONT, bold=True, size=9)
    region_font  = Font(name=_NRB_FONT, bold=True, size=9)
    sub_font     = Font(name=_NRB_FONT, bold=True, size=9)
    grand_font   = Font(name=_NRB_FONT, bold=True, size=10)
    data_font    = Font(name=_NRB_FONT, size=9)

    period = _period_label(year, month)
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731
    COLS = ["LIST OF COUNTIES", "M", "F", "TOTAL", "M", "F", "TOTAL", "GRAND\nTOTAL", "% OF\nTOTAL"]

    # ── Sheet 0: Summary ──
    ws0 = wb.create_sheet("Summary")
    ws0.merge_cells("A1:I1")
    c = ws0.cell(row=1, column=1, value="NATIONAL REGISTRATION BUREAU")
    c.font = title_font
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws0.merge_cells("A2:I2")
    c = ws0.cell(row=2, column=1, value=f"{title} — {period}")
    c.font = subtitle_font
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws0.merge_cells("A4:I4")
    c = ws0.cell(row=4, column=1, value="NARRATIVE SUMMARY")
    c.font = narrative_hdr_font
    ri = 6
    for line in _general_narrative(data, period):
        ws0.merge_cells(f"A{ri}:I{ri}")
        c = ws0.cell(row=ri, column=1, value=line)
        c.font = narrative_font
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        ws0.row_dimensions[ri].height = 45
        ri += 2
    for col_letter in ("A","B","C","D","E","F","G","H","I"):
        ws0.column_dimensions[col_letter].width = 16

    for prefix, module_title in MODULES:
        ws = wb.create_sheet(title=module_title[:31])
        ws.sheet_view.showGridLines = False

        # ── Row 1: Report title ──
        ws.merge_cells("A1:I1")
        c = ws.cell(row=1, column=1, value=f"{module_title} FROM {period}")
        c.font = title_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 24

        # ── Row 2: blank ──
        ws.row_dimensions[2].height = 6

        # ── Row 3: Group headers ──
        ws.merge_cells("B3:D3")
        ws.merge_cells("E3:G3")
        for col, label in [(2, "NPR"), (5, "DUP/OTHERS")]:
            c = ws.cell(row=3, column=col, value=label)
            c.font = hdr_font
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = border
        ws.row_dimensions[3].height = 18

        # ── Row 4: Column headers ──
        for ci, lbl in enumerate(COLS, start=1):
            c = ws.cell(row=4, column=ci, value=lbl)
            c.font  = hdr_font
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border
        ws.row_dimensions[4].height = 28

        # Column widths
        ws.column_dimensions["A"].width = 32
        for col_letter in ["B","C","D","E","F","G","H","I"]:
            ws.column_dimensions[col_letter].width = 10

        # ── Data rows ──
        rows = _module_table_rows(data.get(prefix, {}))
        grand_row = next((r for r in rows if r["_type"] == "grand"), None)
        grand_total = (
            grand_row["npr_m"] + grand_row["npr_f"] + grand_row["oth_m"] + grand_row["oth_f"]
        ) if grand_row else 0

        ri = 5
        for row in rows:
            rtype = row["_type"]

            if rtype == "region":
                ws.merge_cells(start_row=ri, start_column=1, end_row=ri, end_column=9)
                c = ws.cell(row=ri, column=1, value=row["label"])
                c.font   = region_font
                c.fill   = red_fill
                c.border = border
                c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
                ws.row_dimensions[ri].height = 16
                ri += 1
                continue

            is_sub   = rtype == "subtotal"
            is_grand = rtype == "grand"
            fill  = yellow_fill if (is_sub or is_grand) else None
            font_ = grand_font if is_grand else (sub_font if is_sub else data_font)

            npr_m = row.get("npr_m", 0)
            npr_f = row.get("npr_f", 0)
            oth_m = row.get("oth_m", 0)
            oth_f = row.get("oth_f", 0)
            grand = npr_m + npr_f + oth_m + oth_f

            vals = [
                (row["label"], "left"),
                (npr_m,  "right"), (npr_f,  "right"), (npr_m  + npr_f,  "right"),
                (oth_m,  "right"), (oth_f,  "right"), (oth_m  + oth_f,  "right"),
                (grand,  "right"),
                (f"{pct(grand, grand_total):.1f}%", "right"),
            ]
            for ci, (val, align) in enumerate(vals, start=1):
                c = ws.cell(row=ri, column=ci, value=val)
                c.font = font_
                if fill:
                    c.fill = fill
                c.alignment = Alignment(horizontal=align, vertical="center")
                c.border = border
                if isinstance(val, int):
                    c.number_format = "#,##0"
            ws.row_dimensions[ri].height = 15
            ri += 1

        ws.freeze_panes = "A5"

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ── CSV ────────────────────────────────────────────────────────────────────────

def build_csv_report(
    year: int,
    month: "int | None",
    data: dict,
) -> bytes:
    period = _period_label(year, month)
    buf = io.StringIO()
    w   = csv.writer(buf)
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731

    w.writerow(["NARRATIVE SUMMARY"])
    for line in _general_narrative(data, period):
        w.writerow([line])
    w.writerow([])

    for prefix, module_title in MODULES:
        w.writerow([f"{module_title} FROM {period}"])
        w.writerow(["County", "NPR M", "NPR F", "NPR Total",
                    "DUP/Others M", "DUP/Others F", "DUP/Others Total", "Grand Total", "% of Total"])
        rows = _module_table_rows(data.get(prefix, {}))
        grand_row = next((r for r in rows if r["_type"] == "grand"), None)
        grand_total = (
            grand_row["npr_m"] + grand_row["npr_f"] + grand_row["oth_m"] + grand_row["oth_f"]
        ) if grand_row else 0
        for row in rows:
            if row["_type"] == "region":
                w.writerow([row["label"]])
                continue
            nm, nf = row.get("npr_m", 0), row.get("npr_f", 0)
            om, of_ = row.get("oth_m", 0), row.get("oth_f", 0)
            grand = nm + nf + om + of_
            w.writerow([
                row["label"], nm, nf, nm + nf,
                om, of_, om + of_,
                grand, f"{pct(grand, grand_total):.1f}%",
            ])
        w.writerow([])

    return buf.getvalue().encode("utf-8-sig")  # BOM for Excel compatibility


# ── PDF ────────────────────────────────────────────────────────────────────────

def build_pdf_report(
    year: int,
    month: "int | None",
    data: dict,
    org_name: str = "National Registration Bureau",
) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=1.5 * cm,  bottomMargin=1.5 * cm,
    )
    period = _period_label(year, month)
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731

    c_red    = colors.HexColor(f"#{_NRB_RED}")
    c_yellow = colors.HexColor(f"#{_NRB_YELLOW}")
    c_grid   = colors.HexColor("#000000")

    title_style = ParagraphStyle("rptTitle", fontSize=16, fontName="Times-Bold", alignment=1, spaceAfter=4)
    h2_style    = ParagraphStyle("h2", fontSize=11, fontName="Times-Bold", spaceBefore=4, spaceAfter=4)
    body_style  = ParagraphStyle("body", fontSize=10, fontName="Times-Roman", spaceAfter=6, leading=14)

    story = [Paragraph("NATIONAL REGISTRATION BUREAU", title_style)]
    story.append(Paragraph(f"NARRATIVE SUMMARY — {period}", h2_style))
    for line in _general_narrative(data, period):
        story.append(Paragraph(line, body_style))
    story.append(PageBreak())

    for pi, (prefix, module_title) in enumerate(MODULES):
        if pi > 0:
            story.append(PageBreak())

        story.append(Paragraph(f"{module_title} FROM {period}", title_style))
        story.append(Spacer(1, 0.3 * cm))

        rows_data = _module_table_rows(data.get(prefix, {}))
        grand_row = next((r for r in rows_data if r["_type"] == "grand"), None)
        grand_total = (
            grand_row["npr_m"] + grand_row["npr_f"] + grand_row["oth_m"] + grand_row["oth_f"]
        ) if grand_row else 0

        # Table header
        header_row1 = ["LIST OF COUNTIES", "NPR", "", "", "DUP/OTHERS", "", "", "GRAND\nTOTAL", "% OF\nTOTAL"]
        header_row2 = ["", "M", "F", "TOTAL", "M", "F", "TOTAL", "", ""]

        tbl_data = [header_row1, header_row2]
        tbl_style_cmds = [
            ("SPAN", (0, 0), (0, 1)), ("SPAN", (7, 0), (7, 1)), ("SPAN", (8, 0), (8, 1)),
            ("SPAN", (1, 0), (3, 0)), ("SPAN", (4, 0), (6, 0)),
            ("FONTNAME",     (0, 0), (-1, 1), "Times-Bold"),
            ("FONTSIZE",     (0, 0), (-1, 1), 8),
            ("ALIGN",        (0, 0), (-1, 1), "CENTER"),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("FONTNAME",     (0, 2), (-1, -1), "Times-Roman"),
            ("FONTSIZE",     (0, 2), (-1, -1), 8),
            ("ALIGN",        (1, 2), (-1, -1), "RIGHT"),
            ("ALIGN",        (0, 2), (0, -1),  "LEFT"),
            ("GRID",         (0, 0), (-1, -1), 0.4, c_grid),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
        ]

        row_offset = 2
        for i, row in enumerate(rows_data):
            rtype = row["_type"]
            r = row_offset + i

            if rtype == "region":
                tbl_data.append([row["label"], "", "", "", "", "", "", "", ""])
                tbl_style_cmds += [
                    ("SPAN",       (0, r), (-1, r)),
                    ("BACKGROUND", (0, r), (-1, r), c_red),
                    ("FONTNAME",   (0, r), (-1, r), "Times-Bold"),
                ]
                continue

            nm, nf = row.get("npr_m", 0), row.get("npr_f", 0)
            om, of_ = row.get("oth_m", 0), row.get("oth_f", 0)
            grand = nm + nf + om + of_
            tbl_data.append([
                row["label"],
                f"{nm:,}", f"{nf:,}", f"{nm+nf:,}",
                f"{om:,}", f"{of_:,}", f"{om+of_:,}",
                f"{grand:,}", f"{pct(grand, grand_total):.1f}%",
            ])
            if rtype in ("subtotal", "grand"):
                tbl_style_cmds += [
                    ("BACKGROUND", (0, r), (-1, r), c_yellow),
                    ("FONTNAME",   (0, r), (-1, r), "Times-Bold"),
                ]

        pw = landscape(A4)[0] - 3 * cm
        col_widths = [pw * 0.24] + [pw * 0.76 / 8] * 8

        tbl = Table(tbl_data, colWidths=col_widths, repeatRows=2)
        tbl.setStyle(TableStyle(tbl_style_cmds))
        story.append(tbl)

        from datetime import date
        footer = ParagraphStyle("footer", fontSize=8, fontName="Times-Italic", textColor=colors.grey, alignment=1, spaceBefore=8)
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(f"{org_name} · Generated {date.today().strftime('%d %B %Y')}", footer))

    doc.build(story)
    return buffer.getvalue()


# ── Word ───────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, size=7, color_hex=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, font_name=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def build_word_report(
    year: int,
    month: "int | None",
    data: dict,
) -> bytes:
    doc = Document()
    section = doc.sections[0]
    # Landscape
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin  = section.bottom_margin = Inches(0.5)

    period = _period_label(year, month)
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731

    def add_title(text: str, size: int):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = _NRB_FONT

    add_title("NATIONAL REGISTRATION BUREAU", 16)
    add_title(f"NARRATIVE SUMMARY — {period}", 13)
    doc.add_paragraph()
    for line in _general_narrative(data, period):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = _NRB_FONT
    doc.add_page_break()

    for pi, (prefix, module_title) in enumerate(MODULES):
        if pi > 0:
            doc.add_page_break()

        add_title(f"{module_title} FROM {period}", 12)

        # Table: 2 header rows + data rows, plus a % OF TOTAL column
        rows_data = _module_table_rows(data.get(prefix, {}))
        grand_row = next((r for r in rows_data if r["_type"] == "grand"), None)
        grand_total = (
            grand_row["npr_m"] + grand_row["npr_f"] + grand_row["oth_m"] + grand_row["oth_f"]
        ) if grand_row else 0

        nrows = 2 + len(rows_data)
        table = doc.add_table(rows=nrows, cols=9)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row 1: group labels
        h1 = table.rows[0]
        h2 = table.rows[1]
        h1.cells[0].merge(h2.cells[0])
        _cell_text(h1.cells[0], "LIST OF COUNTIES", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)
        h1.cells[1].merge(h1.cells[3])
        _cell_text(h1.cells[1], "NPR", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)
        h1.cells[4].merge(h1.cells[6])
        _cell_text(h1.cells[4], "DUP/OTHERS", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)
        h1.cells[7].merge(h2.cells[7])
        _cell_text(h1.cells[7], "GRAND\nTOTAL", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)
        h1.cells[8].merge(h2.cells[8])
        _cell_text(h1.cells[8], "% OF\nTOTAL", bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)

        # Header row 2: M/F/TOTAL labels
        labels2 = ["", "M", "F", "TOTAL", "M", "F", "TOTAL", "", ""]
        for ci, lbl in enumerate(labels2):
            if lbl:
                _cell_text(h2.cells[ci], lbl, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_NRB_FONT)

        # Data rows
        for ri, row in enumerate(rows_data, start=2):
            trow = table.rows[ri]
            rtype = row["_type"]

            if rtype == "region":
                trow.cells[0].merge(trow.cells[8])
                _cell_text(trow.cells[0], row["label"], bold=True, size=7, font_name=_NRB_FONT)
                _set_cell_bg(trow.cells[0], _NRB_RED)
                continue

            nm, nf = row.get("npr_m", 0), row.get("npr_f", 0)
            om, of_ = row.get("oth_m", 0), row.get("oth_f", 0)
            grand = nm + nf + om + of_
            is_sub   = rtype == "subtotal"
            is_grand = rtype == "grand"
            bg = _NRB_YELLOW if (is_sub or is_grand) else None

            vals = [row["label"], nm, nf, nm+nf, om, of_, om+of_, grand, f"{pct(grand, grand_total):.1f}%"]
            aligns = [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.RIGHT] * 8
            for ci, (val, align) in enumerate(zip(vals, aligns)):
                text = f"{val:,}" if isinstance(val, int) else str(val)
                _cell_text(trow.cells[ci], text, bold=is_sub or is_grand,
                           size=7, align=align, font_name=_NRB_FONT)
                if bg:
                    _set_cell_bg(trow.cells[ci], bg)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ── Usajili Mashinani (mobile registration) reports ─────────────────────────────
# Styled to match NRB's own "Statistics 2016"-era workbook convention (the same
# one stats.xlsx uses for the head-office returns): no fill on headers (plain
# bold black text), a RED row for each county heading, and a YELLOW row for
# each TOTAL/subtotal line. No dark theme, no row striping, no invented accent
# colors — kept to the same plain, recognizable palette NRB already uses.
# `data` is the dict returned by `_mobile_summary_data()` in the reports endpoint:
#   {"breakdown": [...], "county_totals": [...], "totals": {...}, "age_breakdown": [...]}

_MOBILE_FONT   = "Times New Roman"
_NRB_RED    = "FF0000"   # county heading rows
_NRB_YELLOW = "FFFF00"   # TOTAL / GRAND TOTAL rows


def _mobile_period_text(year: int, month: "int | None", quarter: "int | None") -> str:
    if month:
        months = ["JANUARY","FEBRUARY","MARCH","APRIL","MAY","JUNE",
                  "JULY","AUGUST","SEPTEMBER","OCTOBER","NOVEMBER","DECEMBER"]
        return f"{months[month - 1]} {year}"
    if quarter:
        qmonths = {1: "JAN – MAR", 2: "APR – JUN", 3: "JUL – SEP", 4: "OCT – DEC"}
        return f"Q{quarter} {year} ({qmonths[quarter]})"
    return f"FULL YEAR {year} (JAN – DEC)"


def _mobile_narrative(data: dict, period: str) -> list[str]:
    """Plain-English summary paragraph(s) for the period — the same figures
    as the tables, in words, so a director can lift this straight into a
    briefing without re-deriving it from the numbers."""
    t = data["totals"]
    county_totals = data["county_totals"]
    total = t["total_registered"]

    if total == 0:
        return [f"No Usajili Mashinani registrations were recorded for {period}."]

    lines = []
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731
    lines.append(
        f"During {period}, mobile registration (Usajili Mashinani) outreach covered "
        f"{t['counties_covered']} {'county' if t['counties_covered'] == 1 else 'counties'} and "
        f"{t['subcounties_covered']} {'subcounty' if t['subcounties_covered'] == 1 else 'subcounties'}. "
        f"A total of {total:,} people were registered against a combined target of {t['target_set']:,}, "
        f"representing an overall achievement of {t['target_achievement_pct']:.1f}%."
    )
    lines.append(
        f"Of those registered, {t['male_total']:,} ({pct(t['male_total'], total):.1f}%) were male and "
        f"{t['female_total']:,} ({pct(t['female_total'], total):.1f}%) were female. "
        f"Live capture applications accounted for {t['live_total']:,} ({pct(t['live_total'], total):.1f}%) "
        f"of the total, with the remaining {t['manual_total']:,} ({pct(t['manual_total'], total):.1f}%) "
        f"captured manually."
    )
    if county_totals:
        ranked = sorted(county_totals, key=lambda r: r["total_registered"], reverse=True)
        top, bottom = ranked[0], ranked[-1]
        if top["county"] != bottom["county"]:
            lines.append(
                f"{top['county']} County recorded the highest volume at {top['total_registered']:,} "
                f"registered ({top['target_achievement_pct']:.1f}% of its target), while {bottom['county']} "
                f"County recorded the lowest at {bottom['total_registered']:,} "
                f"({bottom['target_achievement_pct']:.1f}% of its target)."
            )
    age_total = t.get("age_grand_total", 0)
    if age_total:
        lines.append(
            f"By age band (Daily Report — NPR), {t['age_25_40_total']:,} ({pct(t['age_25_40_total'], age_total):.1f}%) "
            f"were aged 25–40, {t['age_41_60_total']:,} ({pct(t['age_41_60_total'], age_total):.1f}%) were aged "
            f"41–60, and {t['age_60_plus_total']:,} ({pct(t['age_60_plus_total'], age_total):.1f}%) were above 60 years."
        )
    return lines


def build_mobile_excel_report(data: dict, year: int, month: "int | None", quarter: "int | None") -> bytes:
    wb = Workbook()
    period = _mobile_period_text(year, month, quarter)
    t = data["totals"]

    red_fill    = PatternFill(start_color=_NRB_RED,    end_color=_NRB_RED,    fill_type="solid")
    yellow_fill = PatternFill(start_color=_NRB_YELLOW, end_color=_NRB_YELLOW, fill_type="solid")
    title_font = Font(name=_MOBILE_FONT, bold=True, size=16)
    subtitle_font = Font(name=_MOBILE_FONT, bold=True, size=13)
    period_font = Font(name=_MOBILE_FONT, bold=True, size=13)
    hdr_font   = Font(name=_MOBILE_FONT, bold=True, size=11)
    county_font = Font(name=_MOBILE_FONT, bold=True, size=11)
    total_font  = Font(name=_MOBILE_FONT, bold=True, size=11)
    grand_font = Font(name=_MOBILE_FONT, bold=True, size=12)
    data_font  = Font(name=_MOBILE_FONT, size=11)
    narrative_font = Font(name=_MOBILE_FONT, size=12)
    thin = Side(style="thin", color="000000")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731

    def write_title_block(ws, last_col_letter: str):
        ws.merge_cells(f"A1:{last_col_letter}1")
        c = ws.cell(row=1, column=1, value="NATIONAL REGISTRATION BUREAU")
        c.font = title_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(f"A2:{last_col_letter}2")
        c = ws.cell(row=2, column=1, value="USAJILI MASHINANI — MOBILE REGISTRATION REPORT")
        c.font = subtitle_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.merge_cells(f"A3:{last_col_letter}3")
        c = ws.cell(row=3, column=1, value=period)
        c.font = period_font
        c.alignment = Alignment(horizontal="center", vertical="center")

    def hdr_cell(ws, row, col, value):
        cell = ws.cell(row=row, column=col, value=value)
        cell.font = hdr_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        return cell

    # ── Sheet 0: Summary ──
    ws0 = wb.active
    ws0.title = "Summary"
    write_title_block(ws0, "H")
    ws0.merge_cells("A5:H5")
    c = ws0.cell(row=5, column=1, value="NARRATIVE SUMMARY")
    c.font = Font(name=_MOBILE_FONT, bold=True, size=13)
    ri = 7
    for line in _mobile_narrative(data, period):
        ws0.merge_cells(f"A{ri}:H{ri}")
        c = ws0.cell(row=ri, column=1, value=line)
        c.font = narrative_font
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        ws0.row_dimensions[ri].height = 45
        ri += 2
    for col_letter in ("A","B","C","D","E","F","G","H"):
        ws0.column_dimensions[col_letter].width = 16

    # ── Sheet 1: Target vs Achievement by County ──
    ws1 = wb.create_sheet("Target vs Achievement")
    write_title_block(ws1, "D")

    cols1 = ["County", "Target", "Registered", "Achievement %"]
    for ci, lbl in enumerate(cols1, start=1):
        hdr_cell(ws1, 5, ci, lbl)
    ws1.column_dimensions["A"].width = 24
    for col_letter in ("B", "C", "D"):
        ws1.column_dimensions[col_letter].width = 16

    ri = 6
    for row in data["county_totals"]:
        vals = [row["county"], row["target_set"], row["total_registered"], row["target_achievement_pct"]]
        for ci, val in enumerate(vals, start=1):
            cell = ws1.cell(row=ri, column=ci, value=val)
            cell.font = data_font
            cell.border = border
            cell.alignment = Alignment(horizontal="left" if ci == 1 else "right", vertical="center")
        ri += 1

    vals = ["GRAND TOTAL", t["target_set"], t["total_registered"], t["target_achievement_pct"]]
    for ci, val in enumerate(vals, start=1):
        cell = ws1.cell(row=ri, column=ci, value=val)
        cell.font = grand_font
        cell.fill = yellow_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="left" if ci == 1 else "right", vertical="center")

    # ── Sheet 2: Registration Volume by County, Subcounty & Ward ──
    # Mirrors the field template's own breakdown: Live Capture vs Manual, each
    # split into NPR (Initial) and Replacement ("Duplicate") applications, by
    # M/F — grouped per county with a heading row and a TOTAL row, matching
    # NRB's own county/total row convention (red heading, yellow total).
    ws2 = wb.create_sheet("Registration Volume")
    write_title_block(ws2, "S")

    ws2.merge_cells("A5:A7"); hdr_cell(ws2, 5, 1, "County")
    ws2.merge_cells("B5:B7"); hdr_cell(ws2, 5, 2, "Subcounty")
    ws2.merge_cells("C5:C7"); hdr_cell(ws2, 5, 3, "Ward")
    ws2.merge_cells("D5:J5"); hdr_cell(ws2, 5, 4, "LIVE CAPTURE APPLICATIONS")
    ws2.merge_cells("K5:Q5"); hdr_cell(ws2, 5, 11, "MANUAL APPLICATIONS")
    ws2.merge_cells("R5:R7"); hdr_cell(ws2, 5, 18, "TOTAL\nREGISTERED")
    ws2.merge_cells("S5:S7"); hdr_cell(ws2, 5, 19, "% OF\nTOTAL")

    ws2.merge_cells("D6:F6"); hdr_cell(ws2, 6, 4, "INITIAL (NPR)")
    ws2.merge_cells("G6:I6"); hdr_cell(ws2, 6, 7, "OTHERS (REPLACEMENT)")
    ws2.merge_cells("J6:J7"); hdr_cell(ws2, 6, 10, "SUB\nTOTAL")
    ws2.merge_cells("K6:M6"); hdr_cell(ws2, 6, 11, "INITIAL (NPR)")
    ws2.merge_cells("N6:P6"); hdr_cell(ws2, 6, 14, "OTHERS (REPLACEMENT)")
    ws2.merge_cells("Q6:Q7"); hdr_cell(ws2, 6, 17, "SUB\nTOTAL")

    for col in (4, 7, 11, 14):
        hdr_cell(ws2, 7, col, "M")
        hdr_cell(ws2, 7, col + 1, "F")
        hdr_cell(ws2, 7, col + 2, "T")

    ws2.column_dimensions["A"].width = 18
    ws2.column_dimensions["B"].width = 18
    ws2.column_dimensions["C"].width = 18
    for col_letter in ("D","E","F","G","H","I","J","K","L","M","N","O","P","Q"):
        ws2.column_dimensions[col_letter].width = 7
    ws2.column_dimensions["R"].width = 13
    ws2.column_dimensions["S"].width = 10
    for r in (5, 6, 7):
        ws2.row_dimensions[r].height = 24

    def data_row(ws, row_idx, vals, font, fill=None):
        for ci, val in enumerate(vals, start=1):
            cell = ws.cell(row=row_idx, column=ci, value=val)
            cell.font = font
            cell.border = border
            if fill:
                cell.fill = fill
            cell.alignment = Alignment(horizontal="left" if ci <= 3 else "right", vertical="center")

    def county_heading(ws, row_idx, county, ncols):
        ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=ncols)
        cell = ws.cell(row=row_idx, column=1, value=f"{county.upper()} COUNTY")
        cell.font = county_font
        cell.fill = red_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)

    def _zero_volume_row(label) -> dict:
        return {k: 0 for k in (
            "live_npr_male", "live_npr_female", "live_npr_total",
            "live_replacement_male", "live_replacement_female", "live_replacement_total",
            "manual_npr_male", "manual_npr_female", "manual_npr_total",
            "manual_replacement_male", "manual_replacement_female", "manual_replacement_total",
            "live_total", "manual_total", "total_registered",
        )} | {"county": label}

    ri = 8
    current_county = None
    county_acc = None
    grand_registered = t["total_registered"]

    def flush_county_total(ws, ri):
        nonlocal county_acc
        if county_acc is None:
            return ri
        vals = [
            f"{current_county.upper()} TOTAL", "", "",
            county_acc["live_npr_male"], county_acc["live_npr_female"], county_acc["live_npr_total"],
            county_acc["live_replacement_male"], county_acc["live_replacement_female"], county_acc["live_replacement_total"],
            county_acc["live_total"],
            county_acc["manual_npr_male"], county_acc["manual_npr_female"], county_acc["manual_npr_total"],
            county_acc["manual_replacement_male"], county_acc["manual_replacement_female"], county_acc["manual_replacement_total"],
            county_acc["manual_total"],
            county_acc["total_registered"],
            f"{pct(county_acc['total_registered'], grand_registered):.1f}%",
        ]
        data_row(ws, ri, vals, total_font, fill=yellow_fill)
        return ri + 1

    for row in data["breakdown"]:
        if row["county"] != current_county:
            ri = flush_county_total(ws2, ri)
            current_county = row["county"]
            county_acc = _zero_volume_row(current_county)
            county_heading(ws2, ri, current_county, 19)
            ri += 1
        for k in county_acc:
            if k != "county":
                county_acc[k] += row[k]
        data_row(ws2, ri, [
            row["county"], row["subcounty"], row["ward"],
            row["live_npr_male"], row["live_npr_female"], row["live_npr_total"],
            row["live_replacement_male"], row["live_replacement_female"], row["live_replacement_total"],
            row["live_total"],
            row["manual_npr_male"], row["manual_npr_female"], row["manual_npr_total"],
            row["manual_replacement_male"], row["manual_replacement_female"], row["manual_replacement_total"],
            row["manual_total"],
            row["total_registered"],
            f"{pct(row['total_registered'], grand_registered):.1f}%",
        ], data_font)
        ri += 1
    ri = flush_county_total(ws2, ri)

    data_row(ws2, ri, [
        "GRAND TOTAL", "", "",
        t["live_npr_male"], t["live_npr_female"], t["live_npr_total"],
        t["live_replacement_male"], t["live_replacement_female"], t["live_replacement_total"],
        t["live_total"],
        t["manual_npr_male"], t["manual_npr_female"], t["manual_npr_total"],
        t["manual_replacement_male"], t["manual_replacement_female"], t["manual_replacement_total"],
        t["manual_total"],
        t["total_registered"],
        "100.0%" if grand_registered else "0.0%",
    ], grand_font, fill=yellow_fill)

    # ── Sheet 3: Daily Report — NPR by Age Band ──
    ws3 = wb.create_sheet("NPR Age Band")
    write_title_block(ws3, "K")
    ws3.merge_cells("A5:A6"); hdr_cell(ws3, 5, 1, "County")
    ws3.merge_cells("B5:B6"); hdr_cell(ws3, 5, 2, "Subcounty")
    ws3.merge_cells("C5:E5"); hdr_cell(ws3, 5, 3, "25–40 YEARS")
    ws3.merge_cells("F5:H5"); hdr_cell(ws3, 5, 6, "41–60 YEARS")
    ws3.merge_cells("I5:K5"); hdr_cell(ws3, 5, 9, "ABOVE 60 YEARS")
    ws3.merge_cells("L5:L6"); hdr_cell(ws3, 5, 12, "TOTAL")
    ws3.merge_cells("M5:M6"); hdr_cell(ws3, 5, 13, "% OF\nTOTAL")
    for col in (3, 6, 9):
        hdr_cell(ws3, 6, col, "M")
        hdr_cell(ws3, 6, col + 1, "F")
        hdr_cell(ws3, 6, col + 2, "T")
    ws3.column_dimensions["A"].width = 18
    ws3.column_dimensions["B"].width = 18
    for col_letter in ("C","D","E","F","G","H","I","J","K"):
        ws3.column_dimensions[col_letter].width = 8
    ws3.column_dimensions["L"].width = 11
    ws3.column_dimensions["M"].width = 10
    for r in (5, 6):
        ws3.row_dimensions[r].height = 22

    age_grand = t.get("age_grand_total", 0)
    ri = 7
    current_county = None
    county_acc3 = None

    def flush_age_total(ws, ri):
        nonlocal county_acc3
        if county_acc3 is None or current_county is None:
            return ri
        vals = [
            f"{current_county.upper()} TOTAL", "",
            county_acc3["age_25_40_male"], county_acc3["age_25_40_female"], county_acc3["age_25_40_total"],
            county_acc3["age_41_60_male"], county_acc3["age_41_60_female"], county_acc3["age_41_60_total"],
            county_acc3["age_60_plus_male"], county_acc3["age_60_plus_female"], county_acc3["age_60_plus_total"],
            county_acc3["age_grand_total"],
            f"{pct(county_acc3['age_grand_total'], age_grand):.1f}%",
        ]
        data_row(ws, ri, vals, total_font, fill=yellow_fill)
        return ri + 1

    for row in data.get("age_breakdown", []):
        if row["county"] != current_county:
            ri = flush_age_total(ws3, ri)
            current_county = row["county"]
            county_acc3 = {k: 0 for k in (
                "age_25_40_male", "age_25_40_female", "age_25_40_total",
                "age_41_60_male", "age_41_60_female", "age_41_60_total",
                "age_60_plus_male", "age_60_plus_female", "age_60_plus_total",
                "age_grand_total",
            )}
            county_heading(ws3, ri, current_county, 13)
            ri += 1
        for k in county_acc3:
            county_acc3[k] += row[k]
        data_row(ws3, ri, [
            row["county"], row["subcounty"],
            row["age_25_40_male"], row["age_25_40_female"], row["age_25_40_total"],
            row["age_41_60_male"], row["age_41_60_female"], row["age_41_60_total"],
            row["age_60_plus_male"], row["age_60_plus_female"], row["age_60_plus_total"],
            row["age_grand_total"],
            f"{pct(row['age_grand_total'], age_grand):.1f}%",
        ], data_font)
        ri += 1
    ri = flush_age_total(ws3, ri)

    data_row(ws3, ri, [
        "GRAND TOTAL", "",
        t["age_25_40_male"], t["age_25_40_female"], t["age_25_40_total"],
        t["age_41_60_male"], t["age_41_60_female"], t["age_41_60_total"],
        t["age_60_plus_male"], t["age_60_plus_female"], t["age_60_plus_total"],
        age_grand,
        "100.0%" if age_grand else "0.0%",
    ], grand_font, fill=yellow_fill)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def build_mobile_pdf_report(data: dict, year: int, month: "int | None", quarter: "int | None",
                             org_name: str = "National Registration Bureau") -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=landscape(A4),
        leftMargin=1.5 * cm, rightMargin=1.5 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    period = _mobile_period_text(year, month, quarter)
    c_yellow = colors.HexColor(f"#{_NRB_YELLOW}")
    c_grid = colors.HexColor("#000000")

    title_style    = ParagraphStyle("t1", fontSize=16, fontName="Times-Bold", alignment=1, spaceAfter=2)
    subtitle_style = ParagraphStyle("t2", fontSize=13, fontName="Times-Bold", alignment=1, spaceAfter=2)
    period_style   = ParagraphStyle("t3", fontSize=13, fontName="Times-Bold", alignment=1, spaceAfter=6)
    h2_style       = ParagraphStyle("h2", fontSize=11, fontName="Times-Bold", spaceBefore=4, spaceAfter=4)

    story = [
        Paragraph("NATIONAL REGISTRATION BUREAU", title_style),
        Paragraph("USAJILI MASHINANI — MOBILE REGISTRATION REPORT", subtitle_style),
        Paragraph(period, period_style),
    ]

    # Table 1: Target vs Achievement
    story.append(Paragraph("Target vs. Achievement — by County", h2_style))
    t = data["totals"]
    tbl1 = [["County", "Target", "Registered", "Achievement %"]]
    style1 = [
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"), ("ALIGN", (0, 0), (0, -1), "LEFT"),
        ("GRID", (0, 0), (-1, -1), 0.5, c_grid),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for row in data["county_totals"]:
        tbl1.append([row["county"], f"{row['target_set']:,}", f"{row['total_registered']:,}", f"{row['target_achievement_pct']:.1f}%"])
    grand_row = len(tbl1)
    tbl1.append(["GRAND TOTAL", f"{t['target_set']:,}", f"{t['total_registered']:,}", f"{t['target_achievement_pct']:.1f}%"])
    style1 += [
        ("FONTNAME", (0, grand_row), (-1, grand_row), "Times-Bold"),
        ("BACKGROUND", (0, grand_row), (-1, grand_row), c_yellow),
    ]
    story.append(Table(tbl1, colWidths=[6 * cm, 4 * cm, 4 * cm, 4 * cm], style=TableStyle(style1)))
    story.append(Spacer(1, 0.6 * cm))

    # Table 2: Registration Volume by County, Subcounty & Ward — mirrors the
    # field template's own breakdown: Live Capture vs Manual, each split into
    # NPR (Initial) and Replacement ("Duplicate") applications, by M/F.
    story.append(Paragraph("Registration Volume — by County, Subcounty &amp; Ward", h2_style))
    header_row0 = ["County", "Subcounty", "Ward", "LIVE CAPTURE APPLICATIONS", "", "", "", "", "", "",
                   "MANUAL APPLICATIONS", "", "", "", "", "", "", "TOTAL\nREGISTERED"]
    header_row1 = ["", "", "", "INITIAL (NPR)", "", "", "OTHERS (REPLACEMENT)", "", "", "SUB\nTOTAL",
                   "INITIAL (NPR)", "", "", "OTHERS (REPLACEMENT)", "", "", "SUB\nTOTAL", ""]
    header_row2 = ["", "", "", "M", "F", "T", "M", "F", "T", "", "M", "F", "T", "M", "F", "T", "", ""]
    tbl2 = [header_row0, header_row1, header_row2]
    style2 = [
        ("SPAN", (0, 0), (0, 2)), ("SPAN", (1, 0), (1, 2)), ("SPAN", (2, 0), (2, 2)), ("SPAN", (17, 0), (17, 2)),
        ("SPAN", (3, 0), (9, 0)), ("SPAN", (10, 0), (16, 0)),
        ("SPAN", (3, 1), (5, 1)), ("SPAN", (6, 1), (8, 1)), ("SPAN", (9, 1), (9, 2)),
        ("SPAN", (10, 1), (12, 1)), ("SPAN", (13, 1), (15, 1)), ("SPAN", (16, 1), (16, 2)),
        ("FONTNAME", (0, 0), (-1, 2), "Times-Bold"),
        ("FONTNAME", (0, 3), (-1, -1), "Times-Roman"), ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("ALIGN", (0, 3), (2, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, c_grid),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    for row in data["breakdown"]:
        tbl2.append([
            row["county"], row["subcounty"], row["ward"],
            row["live_npr_male"], row["live_npr_female"], row["live_npr_total"],
            row["live_replacement_male"], row["live_replacement_female"], row["live_replacement_total"],
            row["live_total"],
            row["manual_npr_male"], row["manual_npr_female"], row["manual_npr_total"],
            row["manual_replacement_male"], row["manual_replacement_female"], row["manual_replacement_total"],
            row["manual_total"], row["total_registered"],
        ])
    grand_row2 = len(tbl2)
    tbl2.append([
        "GRAND TOTAL", "", "",
        t["live_npr_male"], t["live_npr_female"], t["live_npr_total"],
        t["live_replacement_male"], t["live_replacement_female"], t["live_replacement_total"],
        t["live_total"],
        t["manual_npr_male"], t["manual_npr_female"], t["manual_npr_total"],
        t["manual_replacement_male"], t["manual_replacement_female"], t["manual_replacement_total"],
        t["manual_total"], t["total_registered"],
    ])
    style2 += [
        ("FONTNAME", (0, grand_row2), (-1, grand_row2), "Times-Bold"),
        ("BACKGROUND", (0, grand_row2), (-1, grand_row2), c_yellow),
        ("FONTSIZE", (0, 0), (-1, 2), 7),  # smaller font in the 3 header rows — labels are long
    ]
    col_w = [2.0 * cm, 2.0 * cm, 2.0 * cm] + [1.0 * cm] * 3 + [1.0 * cm] * 3 + [1.7 * cm] \
        + [1.0 * cm] * 3 + [1.0 * cm] * 3 + [1.7 * cm] + [2.2 * cm]
    story.append(Table(tbl2, colWidths=col_w, style=TableStyle(style2)))

    from datetime import date
    footer = ParagraphStyle("footer", fontSize=8, fontName="Times-Italic", textColor=colors.grey, alignment=1, spaceBefore=10)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(f"{org_name} · Generated {date.today().strftime('%d %B %Y')}", footer))

    doc.build(story)
    return buffer.getvalue()


def build_mobile_word_report(data: dict, year: int, month: "int | None", quarter: "int | None") -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.5)
    section.top_margin = section.bottom_margin = Inches(0.5)

    period = _mobile_period_text(year, month, quarter)

    def add_title(text: str, size: int):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = _MOBILE_FONT

    add_title("NATIONAL REGISTRATION BUREAU", 16)
    add_title("USAJILI MASHINANI — MOBILE REGISTRATION REPORT", 13)
    add_title(period, 13)
    doc.add_paragraph()

    t = data["totals"]
    pct = lambda n, d: round((n / d) * 100, 1) if d else 0.0  # noqa: E731

    doc.add_paragraph().add_run("Narrative Summary").bold = True
    for line in _mobile_narrative(data, period):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = _MOBILE_FONT
    doc.add_paragraph()

    doc.add_paragraph().add_run("Target vs. Achievement — by County").bold = True
    cols1 = ["County", "Target", "Registered", "Achievement %"]
    table1 = doc.add_table(rows=1 + len(data["county_totals"]) + 1, cols=4)
    table1.style = "Table Grid"
    for ci, lbl in enumerate(cols1):
        _cell_text(table1.rows[0].cells[ci], lbl, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
    for ri, row in enumerate(data["county_totals"], start=1):
        vals = [row["county"], f"{row['target_set']:,}", f"{row['total_registered']:,}", f"{row['target_achievement_pct']:.1f}%"]
        for ci, val in enumerate(vals):
            _cell_text(table1.rows[ri].cells[ci], val, size=9, align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT, font_name=_MOBILE_FONT)
    grand_idx = len(data["county_totals"]) + 1
    grand_vals = ["GRAND TOTAL", f"{t['target_set']:,}", f"{t['total_registered']:,}", f"{t['target_achievement_pct']:.1f}%"]
    for ci, val in enumerate(grand_vals):
        _cell_text(table1.rows[grand_idx].cells[ci], val, bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT, font_name=_MOBILE_FONT)
        _set_cell_bg(table1.rows[grand_idx].cells[ci], _NRB_YELLOW)

    doc.add_paragraph()
    doc.add_paragraph().add_run("Registration Volume — by County, Subcounty & Ward").bold = True
    # Mirrors the field template's own breakdown: Live Capture vs Manual, each
    # split into NPR (Initial) and Replacement ("Duplicate") applications, by M/F.
    n_data = len(data["breakdown"])
    grand_registered = t["total_registered"]
    table2 = doc.add_table(rows=3 + n_data + 1, cols=19)
    table2.style = "Table Grid"
    r0, r1, r2 = table2.rows[0], table2.rows[1], table2.rows[2]

    def merge_h(row, c1, c2):
        row.cells[c1].merge(row.cells[c2])

    def merge_v(c, r1_idx, r2_idx):
        table2.rows[r1_idx].cells[c].merge(table2.rows[r2_idx].cells[c])

    merge_v(0, 0, 2); merge_v(1, 0, 2); merge_v(2, 0, 2); merge_v(17, 0, 2); merge_v(18, 0, 2)
    merge_h(r0, 3, 9); merge_h(r0, 10, 16)
    merge_h(r1, 3, 5); merge_h(r1, 6, 8); merge_v(9, 1, 2)
    merge_h(r1, 10, 12); merge_h(r1, 13, 15); merge_v(16, 1, 2)

    for ci, lbl in [(0, "County"), (1, "Subcounty"), (2, "Ward"), (3, "LIVE CAPTURE APPLICATIONS"),
                    (10, "MANUAL APPLICATIONS"), (17, "TOTAL REGISTERED"), (18, "% OF TOTAL")]:
        _cell_text(r0.cells[ci], lbl, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)

    for ci, lbl in [(3, "INITIAL (NPR)"), (6, "OTHERS (REPLACEMENT)"), (9, "SUB TOTAL"),
                    (10, "INITIAL (NPR)"), (13, "OTHERS (REPLACEMENT)"), (16, "SUB TOTAL")]:
        _cell_text(r1.cells[ci], lbl, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)

    for ci in (3, 6, 10, 13):
        _cell_text(r2.cells[ci], "M", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
        _cell_text(r2.cells[ci + 1], "F", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
        _cell_text(r2.cells[ci + 2], "T", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)

    for ri, row in enumerate(data["breakdown"], start=3):
        vals = [
            row["county"], row["subcounty"], row["ward"],
            row["live_npr_male"], row["live_npr_female"], row["live_npr_total"],
            row["live_replacement_male"], row["live_replacement_female"], row["live_replacement_total"],
            row["live_total"],
            row["manual_npr_male"], row["manual_npr_female"], row["manual_npr_total"],
            row["manual_replacement_male"], row["manual_replacement_female"], row["manual_replacement_total"],
            row["manual_total"], row["total_registered"],
            f"{pct(row['total_registered'], grand_registered):.1f}%",
        ]
        for ci, val in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci <= 2 else WD_ALIGN_PARAGRAPH.RIGHT
            _cell_text(table2.rows[ri].cells[ci], f"{val:,}" if isinstance(val, int) else val, size=8, align=align, font_name=_MOBILE_FONT)

    grand_idx2 = 3 + n_data
    grand_vals2 = [
        "GRAND TOTAL", "", "",
        t["live_npr_male"], t["live_npr_female"], t["live_npr_total"],
        t["live_replacement_male"], t["live_replacement_female"], t["live_replacement_total"],
        t["live_total"],
        t["manual_npr_male"], t["manual_npr_female"], t["manual_npr_total"],
        t["manual_replacement_male"], t["manual_replacement_female"], t["manual_replacement_total"],
        t["manual_total"], t["total_registered"],
        "100.0%" if grand_registered else "0.0%",
    ]
    for ci, val in enumerate(grand_vals2):
        align = WD_ALIGN_PARAGRAPH.LEFT if ci <= 2 else WD_ALIGN_PARAGRAPH.RIGHT
        text = f"{val:,}" if isinstance(val, int) else val
        _cell_text(table2.rows[grand_idx2].cells[ci], text, bold=True, size=8, align=align, font_name=_MOBILE_FONT)
        _set_cell_bg(table2.rows[grand_idx2].cells[ci], _NRB_YELLOW)

    # ── Table 3: Daily Report — NPR by Age Band ──
    doc.add_paragraph()
    doc.add_paragraph().add_run("Daily Report — NPR by Age Band").bold = True
    age_breakdown = data.get("age_breakdown", [])
    age_grand = t.get("age_grand_total", 0)
    n_age = len(age_breakdown)
    table3 = doc.add_table(rows=2 + n_age + 1, cols=13)
    table3.style = "Table Grid"
    a0, a1 = table3.rows[0], table3.rows[1]

    def merge_h3(row, c1, c2):
        row.cells[c1].merge(row.cells[c2])

    def merge_v3(c, r1_idx, r2_idx):
        table3.rows[r1_idx].cells[c].merge(table3.rows[r2_idx].cells[c])

    merge_v3(0, 0, 1); merge_v3(1, 0, 1); merge_v3(11, 0, 1); merge_v3(12, 0, 1)
    merge_h3(a0, 2, 4); merge_h3(a0, 5, 7); merge_h3(a0, 8, 10)
    for ci, lbl in [(0, "County"), (1, "Subcounty"), (2, "25–40 YEARS"), (5, "41–60 YEARS"),
                    (8, "ABOVE 60 YEARS"), (11, "TOTAL"), (12, "% OF TOTAL")]:
        _cell_text(a0.cells[ci], lbl, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
    for col in (2, 5, 8):
        _cell_text(a1.cells[col], "M", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
        _cell_text(a1.cells[col + 1], "F", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)
        _cell_text(a1.cells[col + 2], "T", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=_MOBILE_FONT)

    for ri, row in enumerate(age_breakdown, start=2):
        vals = [
            row["county"], row["subcounty"],
            row["age_25_40_male"], row["age_25_40_female"], row["age_25_40_total"],
            row["age_41_60_male"], row["age_41_60_female"], row["age_41_60_total"],
            row["age_60_plus_male"], row["age_60_plus_female"], row["age_60_plus_total"],
            row["age_grand_total"],
            f"{pct(row['age_grand_total'], age_grand):.1f}%",
        ]
        for ci, val in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
            _cell_text(table3.rows[ri].cells[ci], f"{val:,}" if isinstance(val, int) else val, size=8, align=align, font_name=_MOBILE_FONT)

    grand_idx3 = 2 + n_age
    grand_vals3 = [
        "GRAND TOTAL", "",
        t["age_25_40_male"], t["age_25_40_female"], t["age_25_40_total"],
        t["age_41_60_male"], t["age_41_60_female"], t["age_41_60_total"],
        t["age_60_plus_male"], t["age_60_plus_female"], t["age_60_plus_total"],
        age_grand,
        "100.0%" if age_grand else "0.0%",
    ]
    for ci, val in enumerate(grand_vals3):
        align = WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
        text = f"{val:,}" if isinstance(val, int) else val
        _cell_text(table3.rows[grand_idx3].cells[ci], text, bold=True, size=8, align=align, font_name=_MOBILE_FONT)
        _set_cell_bg(table3.rows[grand_idx3].cells[ci], _NRB_YELLOW)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
